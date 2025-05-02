from pathlib import Path
from typing import Dict, Any, List, Union
import os
from docx import Document as DocxDocument
import win32com.client
import pythoncom
import tempfile
import shutil

from autotask.document.base import Document
from autotask.document.reader.file_reader import FileReader
from autotask.document.reader.reader_registry import ReaderRegistry
from autotask.utils.log import logger

CHONKIE_AVAILABLE = None

def _get_chunker():
    global CHONKIE_AVAILABLE
    if CHONKIE_AVAILABLE is None:
        try:
            from chonkie import RecursiveChunker, RecursiveRules
            CHONKIE_AVAILABLE = True
            return RecursiveChunker, RecursiveRules
        except ImportError:
            CHONKIE_AVAILABLE = False
            logger.warning("Chonkie library not installed. Smart chunking will be disabled. Install with: pip install chonkie")
            return None, None
    elif CHONKIE_AVAILABLE:
        from chonkie import RecursiveChunker, RecursiveRules
        return RecursiveChunker, RecursiveRules
    else:
        return None, None


class BaseWordReader(FileReader):
    """Base class for Word document readers with common functionality"""
    
    PARAMS = {
        "chunk_size": {
            "label": "Chunk Size",
            "description": "Maximum tokens per chunk",
            "type": "INTEGER",
            "default": 1024
        },
        "extract_metadata": {
            "label": "Extract Metadata",
            "description": "Whether to extract Word document metadata",
            "type": "BOOLEAN",
            "default": True
        },
        "include_headers": {
            "label": "Include Headers",
            "description": "Whether to include headers in extracted text",
            "type": "BOOLEAN",
            "default": True
        }
    }

    def _create_documents(self, content: str, file_path: Path, metadata: Dict[str, Any], chunk_size: int) -> List[Document]:
        """Create Document objects from content with optional chunking
        
        Args:
            content: Document content text
            file_path: Source file path
            metadata: Document metadata
            chunk_size: Maximum chunk size in tokens
            
        Returns:
            List[Document]: List of document chunks
        """
        RecursiveChunker, RecursiveRules = _get_chunker()
        if len(content) < 1000 or not RecursiveChunker:
            document = Document(
                id=str(file_path),
                name=file_path.stem,
                content=content,
                meta_data=metadata
            )
            return [document]
        
        try:
            logger.info(f"Applying smart chunking to Word document: {file_path}")
            
            chunker = RecursiveChunker(
                tokenizer_or_token_counter="gpt2",
                chunk_size=chunk_size,
                rules=RecursiveRules(),
                min_characters_per_chunk=12,
                return_type="chunks"
            )
            
            chunks = chunker(content)
            logger.info(f"Split Word document into {len(chunks)} chunks")
            
            documents = []
            for i, chunk in enumerate(chunks):
                chunk_metadata = metadata.copy()
                chunk_metadata["chunk_id"] = i + 1
                
                chunk_doc = Document(
                    id=f"{str(file_path)}_{i+1}",
                    name=f"{file_path.stem}_chunk_{i+1}",
                    content=chunk.text,
                    meta_data=chunk_metadata
                )
                documents.append(chunk_doc)
            
            return documents
            
        except Exception as chunk_error:
            logger.error(f"Smart chunking failed: {str(chunk_error)}, returning full document")
            document = Document(
                id=str(file_path),
                name=file_path.stem,
                content=content,
                meta_data=metadata
            )
            return [document]


@ReaderRegistry.register_reader([".docx"])
class DocxReader(BaseWordReader):
    """Reader for .docx format Word documents"""
    
    FILE_FORMATS = [".docx"]
    
    def read_file(self, file_path: Path, params: Dict[str, Any]) -> List[Document]:
        """Read .docx document and return list of Document objects
        
        Args:
            file_path: Path to .docx document
            params: Reader parameters
            
        Returns:
            List[Document]: List of documents with Word document content
        """
        chunk_size = params.get("chunk_size", 1024)
        extract_metadata = params.get("extract_metadata", True)
        include_headers = params.get("include_headers", True)
        
        try:
            # Read Word document
            try:
                doc = DocxDocument(file_path)
            except Exception as e:
                logger.error(f"Failed to open document {file_path}: {str(e)}")
                raise ValueError(f"Invalid or corrupted .docx document: {file_path}") from e
            
            # Extract text from all paragraphs
            content = ""
            for para in doc.paragraphs:
                if para.text.strip():
                    content += para.text + "\n\n"
            
            # Include headers if requested
            if include_headers:
                for section in doc.sections:
                    header = section.header
                    if header.text.strip():
                        content = header.text + "\n\n" + content
            
            # Create metadata
            metadata = {
                "source": str(file_path),
                "filename": file_path.name,
                "extension": file_path.suffix,
                "size": os.path.getsize(file_path),
                "modified": os.path.getmtime(file_path)
            }
            
            # Add Word-specific metadata if requested
            if extract_metadata:
                core_props = doc.core_properties
                if core_props:
                    metadata.update({
                        "title": core_props.title or "",
                        "author": core_props.author or "",
                        "subject": core_props.subject or "",
                        "created": str(core_props.created) if core_props.created else "",
                        "modified": str(core_props.modified) if core_props.modified else "",
                        "last_modified_by": core_props.last_modified_by or "",
                        "revision": core_props.revision or 0
                    })
            
            return self._create_documents(content, file_path, metadata, chunk_size)
            
        except Exception as e:
            logger.error(f"Failed to read .docx document {file_path}: {str(e)}")
            raise


@ReaderRegistry.register_reader([".doc"])
class DocReader(BaseWordReader):
    """
    Reader for .doc format Word documents
    
    Requires:
    - Windows OS
    - Microsoft Word installed
    - pywin32 package installed (pip install pywin32)
    """
    
    FILE_FORMATS = [".doc"]
    
    def read_file(self, file_path: Path, params: Dict[str, Any]) -> List[Document]:
        """Read .doc document and return list of Document objects
        
        Args:
            file_path: Path to .doc document
            params: Reader parameters
            
        Returns:
            List[Document]: List of documents with Word document content
        """
        chunk_size = params.get("chunk_size", 1024)
        extract_metadata = params.get("extract_metadata", True)
        include_headers = params.get("include_headers", True)
        
        pythoncom.CoInitialize()
        word = None
        
        try:
            # Create Word application instance
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            # Open document
            doc = word.Documents.Open(str(file_path.absolute()))
            
            try:
                # Extract text content
                content = ""
                for para in doc.Paragraphs:
                    if para.Range.Text.strip():
                        content += para.Range.Text.strip() + "\n\n"
                
                # Include headers if requested
                if include_headers:
                    for section in range(1, doc.Sections.Count + 1):
                        header_text = doc.Sections(section).Headers(1).Range.Text
                        if header_text.strip():
                            content = header_text.strip() + "\n\n" + content
                
                # Create metadata
                metadata = {
                    "source": str(file_path),
                    "filename": file_path.name,
                    "extension": file_path.suffix,
                    "size": os.path.getsize(file_path),
                    "modified": os.path.getmtime(file_path)
                }
                
                # Add Word-specific metadata if requested
                if extract_metadata:
                    metadata.update({
                        "title": doc.BuiltInDocumentProperties("Title").Value if doc.BuiltInDocumentProperties("Title") else "",
                        "author": doc.BuiltInDocumentProperties("Author").Value if doc.BuiltInDocumentProperties("Author") else "",
                        "subject": doc.BuiltInDocumentProperties("Subject").Value if doc.BuiltInDocumentProperties("Subject") else "",
                        "created": str(doc.BuiltInDocumentProperties("Creation Date").Value) if doc.BuiltInDocumentProperties("Creation Date") else "",
                        "modified": str(doc.BuiltInDocumentProperties("Last Save Time").Value) if doc.BuiltInDocumentProperties("Last Save Time") else "",
                        "last_modified_by": doc.BuiltInDocumentProperties("Last Author").Value if doc.BuiltInDocumentProperties("Last Author") else "",
                        "revision": doc.BuiltInDocumentProperties("Revision Number").Value if doc.BuiltInDocumentProperties("Revision Number") else 0
                    })
                
                doc.Close()
                return self._create_documents(content, file_path, metadata, chunk_size)
                
            finally:
                if doc:
                    try:
                        doc.Close()
                    except:
                        pass
                        
        except Exception as e:
            logger.error(f"Failed to read .doc document {file_path}: {str(e)}")
            raise
            
        finally:
            if word:
                try:
                    word.Quit()
                except:
                    pass
                    
            try:
                pythoncom.CoUninitialize()
            except:
                pass 